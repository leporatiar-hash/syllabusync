from fastapi import FastAPI, UploadFile, File, HTTPException, Query, Request, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel
from dotenv import load_dotenv
from openai import OpenAI, AuthenticationError as OAIAuthError, RateLimitError as OAIRateLimitError, APIStatusError as OAIAPIStatusError
from jose import JWTError, jwt, jwk
from jose.utils import base64url_decode
from sqlalchemy import create_engine, Column, String, Boolean, Date, DateTime, ForeignKey, Text, JSON, Integer, text
from sqlalchemy.dialects.postgresql import UUID as PG_UUID
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker, relationship
from datetime import datetime, date, timedelta
from urllib.request import urlopen
from urllib.error import URLError
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded
import PyPDF2
from docx import Document
import json
import uuid
import io
import os
import re
import base64
import time
import asyncio
import logging
import resend

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()

# Supabase auth configuration
SUPABASE_URL = os.getenv("SUPABASE_URL") or os.getenv("NEXT_PUBLIC_SUPABASE_URL") or ""
SUPABASE_JWKS_URL = os.getenv("SUPABASE_JWKS_URL") or (
    f"{SUPABASE_URL.rstrip('/')}/auth/v1/.well-known/jwks.json" if SUPABASE_URL else ""
)
SUPABASE_ISSUER = os.getenv("SUPABASE_ISSUER") or (
    f"{SUPABASE_URL.rstrip('/')}/auth/v1" if SUPABASE_URL else ""
)
SUPABASE_JWT_AUD = os.getenv("SUPABASE_JWT_AUD", "authenticated")
SUPABASE_JWKS_CACHE_TTL = int(os.getenv("SUPABASE_JWKS_CACHE_TTL", "3600"))

# Log Supabase auth configuration at startup
print(f"[Auth] SUPABASE_URL={'SET (' + SUPABASE_URL[:40] + '...)' if SUPABASE_URL else 'NOT SET'}")
print(f"[Auth] SUPABASE_JWKS_URL={'SET' if SUPABASE_JWKS_URL else 'NOT SET'}")
print(f"[Auth] SUPABASE_ISSUER={'SET' if SUPABASE_ISSUER else 'NOT SET'}")

bearer_scheme = HTTPBearer(auto_error=True)

app = FastAPI()

# Rate limiting setup
limiter = Limiter(key_func=get_remote_address)
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

# Parse ALLOWED_ORIGINS - strip whitespace and quotes from each origin
# Railway/Vercel env vars sometimes include literal quotes that need to be removed
allowed_origins_str = os.getenv("ALLOWED_ORIGINS", "http://localhost:3000")
# Strip outer quotes from the entire string first (in case it's wrapped in quotes)
allowed_origins_str = allowed_origins_str.strip().strip('"').strip("'")
ALLOWED_ORIGINS = [origin.strip().strip('"').strip("'") for origin in allowed_origins_str.split(",") if origin.strip()]

# Log CORS configuration at module load
logger.info(f"[CORS] Configured origins: {ALLOWED_ORIGINS}")

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS", "PATCH"],
    allow_headers=["*"],
)


# Request timeout middleware
@app.middleware("http")
async def timeout_middleware(request: Request, call_next):
    try:
        return await asyncio.wait_for(call_next(request), timeout=60.0)
    except asyncio.TimeoutError:
        return JSONResponse({"detail": "Request timeout"}, status_code=504)


# Initialize OpenAI client
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
RESEND_API_KEY = os.getenv("RESEND_API_KEY", "")
FEEDBACK_EMAIL = os.getenv("FEEDBACK_EMAIL", "")


def _raise_if_openai_error(e: Exception) -> None:
    """If *e* is a known OpenAI SDK error, raise an HTTPException with a
    clear, user-facing message and log it so Railway picks it up.
    For anything else, do nothing — let the caller handle it."""
    if isinstance(e, OAIAuthError):
        logger.error("[OpenAI] Authentication error — API key is invalid or revoked: %s", e)
        raise HTTPException(status_code=503, detail="AI service is temporarily unavailable. Please try again later.")
    if isinstance(e, OAIRateLimitError):
        logger.warning("[OpenAI] Rate-limit hit: %s", e)
        raise HTTPException(status_code=503, detail="AI service is busy right now. Please wait a moment and try again.")
    if isinstance(e, OAIAPIStatusError):
        logger.error("[OpenAI] API error (status %s): %s", e.status_code, e)
        raise HTTPException(status_code=503, detail="AI service returned an error. Please try again later.")


# Supabase auth helpers
_jwks_cache: dict[str, object] = {"fetched_at": 0.0, "keys": {}, "last_fetch_error": None}


def _fetch_jwks(retries: int = 2) -> dict[str, dict] | None:
    """Fetch JWKS from Supabase. Returns None on failure instead of raising."""
    if not SUPABASE_JWKS_URL:
        print("[Auth] ERROR: SUPABASE_JWKS_URL is not configured")
        _jwks_cache["last_fetch_error"] = "JWKS URL not configured"
        return None

    last_error = None
    for attempt in range(retries):
        try:
            print(f"[Auth] Fetching JWKS from {SUPABASE_JWKS_URL} (attempt {attempt + 1}/{retries})")
            with urlopen(SUPABASE_JWKS_URL, timeout=10) as resp:
                payload = json.load(resp)
            keys = {key.get("kid"): key for key in payload.get("keys", []) if key.get("kid")}
            if not keys:
                print("[Auth] WARNING: JWKS response contained no valid keys")
                _jwks_cache["last_fetch_error"] = "JWKS empty"
                return None
            _jwks_cache["fetched_at"] = time.time()
            _jwks_cache["keys"] = keys
            _jwks_cache["last_fetch_error"] = None
            print(f"[Auth] Successfully fetched {len(keys)} keys from JWKS")
            return keys
        except URLError as e:
            last_error = str(e)
            print(f"[Auth] JWKS fetch attempt {attempt + 1} failed: {e}")
            if attempt < retries - 1:
                time.sleep(0.5)  # Brief wait before retry
        except json.JSONDecodeError as e:
            print(f"[Auth] ERROR: Failed to parse JWKS response as JSON: {e}")
            _jwks_cache["last_fetch_error"] = f"JSON parse error: {e}"
            return None

    print(f"[Auth] WARNING: All {retries} JWKS fetch attempts failed. Last error: {last_error}")
    _jwks_cache["last_fetch_error"] = last_error
    return None


def _get_jwks() -> dict[str, dict]:
    """Get JWKS, using stale cache if fresh fetch fails."""
    cached_keys = _jwks_cache.get("keys", {})
    fetched_at = float(_jwks_cache.get("fetched_at", 0.0) or 0.0)
    cache_expired = (time.time() - fetched_at) > SUPABASE_JWKS_CACHE_TTL

    # If cache is fresh, use it
    if cached_keys and not cache_expired:
        return cached_keys  # type: ignore[return-value]

    # Try to fetch fresh keys
    fresh_keys = _fetch_jwks()
    if fresh_keys:
        return fresh_keys

    # Fetch failed - use stale cache if available
    if cached_keys:
        print("[Auth] WARNING: Using stale JWKS cache due to fetch failure")
        return cached_keys  # type: ignore[return-value]

    # No cache and fetch failed - return empty (will cause auth to fail gracefully)
    return {}


def _verify_supabase_token(token: str) -> dict:
    if not SUPABASE_ISSUER or not SUPABASE_JWKS_URL:
        print("[Auth] ERROR: Supabase auth not configured - ISSUER or JWKS_URL missing")
        raise HTTPException(status_code=500, detail="Supabase auth not configured. Check SUPABASE_URL env var.")
    try:
        header = jwt.get_unverified_header(token)
    except JWTError:
        raise HTTPException(status_code=401, detail="Invalid token header")
    kid = header.get("kid")
    if not kid:
        raise HTTPException(status_code=401, detail="Invalid token - missing key ID")

    keys = _get_jwks()
    if not keys:
        # No keys available at all - this is an auth infrastructure issue
        last_error = _jwks_cache.get("last_fetch_error", "Unknown")
        print(f"[Auth] ERROR: No JWKS keys available. Last fetch error: {last_error}")
        raise HTTPException(status_code=503, detail=f"Auth service temporarily unavailable: {last_error}")

    key_data = keys.get(kid)
    if not key_data:
        # Key not in cache - try one more fresh fetch
        fresh_keys = _fetch_jwks()
        if fresh_keys:
            key_data = fresh_keys.get(kid)
    if not key_data:
        raise HTTPException(status_code=401, detail="Invalid token - unknown key ID")
    public_key = jwk.construct(key_data)
    message, encoded_sig = token.rsplit(".", 1)
    decoded_sig = base64url_decode(encoded_sig.encode("utf-8"))
    if not public_key.verify(message.encode("utf-8"), decoded_sig):
        raise HTTPException(status_code=401, detail="Invalid token signature")
    try:
        return jwt.decode(
            token,
            key_data,
            algorithms=[header.get("alg", "RS256")],
            audience=SUPABASE_JWT_AUD,
            issuer=SUPABASE_ISSUER,
        )
    except JWTError:
        raise HTTPException(status_code=401, detail="Invalid token")

# Database setup
DATABASE_URL = os.getenv("DATABASE_URL", "sqlite:///./railway.db")
# Handle Railway's postgres:// vs postgresql://
if DATABASE_URL.startswith("postgres://"):
    DATABASE_URL = DATABASE_URL.replace("postgres://", "postgresql://", 1)

# Use String for UUID to support both SQLite and Postgres
connect_args = {"check_same_thread": False} if DATABASE_URL.startswith("sqlite") else {}
def _safe_db_url(url: str) -> str:
    if "://" not in url:
        return url
    scheme, rest = url.split("://", 1)
    if "@" in rest:
        return f"{scheme}://****:****@{rest.split('@', 1)[1]}"
    return url


print(f"[DB] Using DATABASE_URL={_safe_db_url(DATABASE_URL)}")

engine = create_engine(DATABASE_URL, connect_args=connect_args)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()


def generate_uuid():
    return str(uuid.uuid4())


def ensure_user_columns():
    if engine.dialect.name != "sqlite":
        return
    tables = [
        "courses",
        "deadlines",
        "flashcard_sets",
        "flashcards",
        "quizzes",
        "quiz_questions",
        "summaries",
        "calendar_entries",
        "user_profiles",
    ]
    with engine.begin() as conn:
        for table in tables:
            cols = [row[1] for row in conn.execute(text(f"PRAGMA table_info({table})"))]
            if not cols:
                continue
            if "user_id" not in cols:
                conn.execute(text(f"ALTER TABLE {table} ADD COLUMN user_id TEXT"))

        course_cols = [row[1] for row in conn.execute(text("PRAGMA table_info(courses)"))]
        if course_cols and "course_info" not in course_cols:
            conn.execute(text("ALTER TABLE courses ADD COLUMN course_info TEXT"))

        profile_cols = [row[1] for row in conn.execute(text("PRAGMA table_info(user_profiles)"))]
        if profile_cols:
            if "profile_picture" not in profile_cols:
                conn.execute(text("ALTER TABLE user_profiles ADD COLUMN profile_picture TEXT"))
            if "email" not in profile_cols:
                conn.execute(text("ALTER TABLE user_profiles ADD COLUMN email TEXT"))
            # Backfill defaults for legacy rows
            conn.execute(text(
                "UPDATE user_profiles SET email = user_id || '@legacy.local' WHERE email IS NULL"
            ))

        # Backfill user_id from parent relationships when possible
        conn.execute(text("UPDATE courses SET user_id = COALESCE(user_id, 'legacy') WHERE user_id IS NULL"))
        conn.execute(text(
            "UPDATE deadlines SET user_id = (SELECT user_id FROM courses WHERE courses.id = deadlines.course_id) "
            "WHERE user_id IS NULL"
        ))
        conn.execute(text(
            "UPDATE flashcard_sets SET user_id = (SELECT user_id FROM courses WHERE courses.id = flashcard_sets.course_id) "
            "WHERE user_id IS NULL"
        ))
        conn.execute(text(
            "UPDATE flashcards SET user_id = (SELECT user_id FROM flashcard_sets WHERE flashcard_sets.id = flashcards.flashcard_set_id) "
            "WHERE user_id IS NULL"
        ))
        conn.execute(text(
            "UPDATE summaries SET user_id = (SELECT user_id FROM courses WHERE courses.id = summaries.course_id) "
            "WHERE user_id IS NULL"
        ))
        conn.execute(text(
            "UPDATE quizzes SET user_id = (SELECT user_id FROM courses WHERE courses.id = quizzes.course_id) "
            "WHERE user_id IS NULL"
        ))
        conn.execute(text(
            "UPDATE quiz_questions SET user_id = (SELECT user_id FROM quizzes WHERE quizzes.id = quiz_questions.quiz_id) "
            "WHERE user_id IS NULL"
        ))
        conn.execute(text(
            "UPDATE calendar_entries SET user_id = (SELECT user_id FROM deadlines WHERE deadlines.id = calendar_entries.deadline_id) "
            "WHERE user_id IS NULL"
        ))


# Database Models
class Course(Base):
    __tablename__ = "courses"

    id = Column(String, primary_key=True, default=generate_uuid)
    user_id = Column(String, nullable=False)
    name = Column(String, nullable=False)
    code = Column(String)  # e.g., "FINC 313"
    semester = Column(String)
    start_date = Column(Date)
    end_date = Column(Date)
    course_info = Column(JSON, nullable=True)  # Stores extracted syllabus details (instructor, policies, etc.)
    created_at = Column(DateTime, default=datetime.utcnow)

    deadlines = relationship("Deadline", back_populates="course", cascade="all, delete-orphan")
    flashcard_sets = relationship("FlashcardSet", back_populates="course", cascade="all, delete-orphan")
    quizzes = relationship("Quiz", back_populates="course", cascade="all, delete-orphan")
    summaries = relationship("Summary", back_populates="course", cascade="all, delete-orphan")


class Deadline(Base):
    __tablename__ = "deadlines"

    id = Column(String, primary_key=True, default=generate_uuid)
    user_id = Column(String, nullable=False)
    course_id = Column(String, ForeignKey("courses.id"), nullable=False)
    date = Column(String)  # YYYY-MM-DD
    time = Column(String)  # 11:59pm, etc.
    type = Column(String)  # exam, assignment, quiz, project, reading, deadline
    title = Column(String)
    description = Column(Text)
    recurring = Column(Boolean, default=False)
    frequency = Column(String)  # weekly, biweekly, monthly, null
    day_of_week = Column(String)  # Monday, Tuesday, etc.
    completed = Column(Boolean, default=False)
    created_at = Column(DateTime, default=datetime.utcnow)

    course = relationship("Course", back_populates="deadlines")


class FlashcardSet(Base):
    __tablename__ = "flashcard_sets"

    id = Column(String, primary_key=True, default=generate_uuid)
    user_id = Column(String, nullable=False)
    course_id = Column(String, ForeignKey("courses.id"), nullable=False)
    name = Column(String, nullable=False)
    created_at = Column(DateTime, default=datetime.utcnow)

    course = relationship("Course", back_populates="flashcard_sets")
    flashcards = relationship("Flashcard", back_populates="flashcard_set", cascade="all, delete-orphan")


class Flashcard(Base):
    __tablename__ = "flashcards"

    id = Column(String, primary_key=True, default=generate_uuid)
    user_id = Column(String, nullable=False)
    flashcard_set_id = Column(String, ForeignKey("flashcard_sets.id"), nullable=False)
    front = Column(Text, nullable=False)
    back = Column(Text, nullable=False)
    created_at = Column(DateTime, default=datetime.utcnow)

    flashcard_set = relationship("FlashcardSet", back_populates="flashcards")


class Summary(Base):
    __tablename__ = "summaries"

    id = Column(String, primary_key=True, default=generate_uuid)
    user_id = Column(String, nullable=False)
    course_id = Column(String, ForeignKey("courses.id"), nullable=False)
    title = Column(String, nullable=False)
    content = Column(Text, nullable=False)
    created_at = Column(DateTime, default=datetime.utcnow)

    course = relationship("Course", back_populates="summaries")


class User(Base):
    __tablename__ = "users"

    id = Column(String, primary_key=True, default=generate_uuid)
    email = Column(String, unique=True, nullable=False)
    created_at = Column(DateTime, default=datetime.utcnow)


class UserProfile(Base):
    __tablename__ = "user_profiles"

    id = Column(String, primary_key=True, default=generate_uuid)
    user_id = Column(String, unique=True, nullable=False)
    email = Column(String, unique=True, nullable=False)
    full_name = Column(String, nullable=True)
    school_name = Column(String, nullable=True)
    school_type = Column(String, nullable=True)
    academic_year = Column(String, nullable=True)
    major = Column(String, nullable=True)
    profile_picture = Column(Text, nullable=True)
    created_at = Column(DateTime, default=datetime.utcnow)


class CalendarEntry(Base):
    """Tracks which deadlines have been saved to the user's calendar."""
    __tablename__ = "calendar_entries"

    id = Column(String, primary_key=True, default=generate_uuid)
    deadline_id = Column(String, ForeignKey("deadlines.id"), nullable=False, unique=True)
    user_id = Column(String, nullable=False)
    created_at = Column(DateTime, default=datetime.utcnow)

    deadline = relationship("Deadline", backref="calendar_entry")


class Quiz(Base):
    """A quiz generated from study materials."""
    __tablename__ = "quizzes"

    id = Column(String, primary_key=True, default=generate_uuid)
    user_id = Column(String, nullable=False)
    course_id = Column(String, ForeignKey("courses.id"), nullable=False)
    name = Column(String, nullable=False)
    created_at = Column(DateTime, default=datetime.utcnow)

    course = relationship("Course", back_populates="quizzes")
    questions = relationship("QuizQuestion", back_populates="quiz", cascade="all, delete-orphan")


class QuizQuestion(Base):
    """A single question in a quiz."""
    __tablename__ = "quiz_questions"

    id = Column(String, primary_key=True, default=generate_uuid)
    user_id = Column(String, nullable=False)
    quiz_id = Column(String, ForeignKey("quizzes.id"), nullable=False)
    question = Column(Text, nullable=False)
    options = Column(Text, nullable=False)  # JSON array stored as text
    correct_answer = Column(String, nullable=False)  # A, B, C, or D
    explanation = Column(Text)
    order_num = Column(String, default="0")

    quiz = relationship("Quiz", back_populates="questions")


class Feedback(Base):
    __tablename__ = "feedback"

    id = Column(String, primary_key=True, default=generate_uuid)
    user_email = Column(String, nullable=False)
    feedback_type = Column(String, nullable=False)
    message = Column(Text, nullable=False)
    created_at = Column(DateTime, default=datetime.utcnow)


class ProfilePictureRequest(BaseModel):
    image_data: str  # base64-encoded image

class CreateCourseRequest(BaseModel):
    name: str
    code: str | None = None
    semester: str | None = None
    start_date: str | None = None
    end_date: str | None = None


class UpdateCourseRequest(BaseModel):
    name: str | None = None
    code: str | None = None
    semester: str | None = None
    start_date: str | None = None
    end_date: str | None = None
    course_info: dict | None = None


class UpdateDeadlineRequest(BaseModel):
    date: str


class UserProfileRequest(BaseModel):
    email: str | None = None
    full_name: str | None = None
    school_name: str | None = None
    school_type: str | None = None
    academic_year: str | None = None
    major: str | None = None


class CreateDeadlineRequest(BaseModel):
    title: str
    date: str
    time: str | None = None
    type: str = "Deadline"
    description: str | None = None
    course_id: str | None = None


class QuizSubmission(BaseModel):
    answers: dict[str, str]  # question_id -> selected answer (A, B, C, D)


class FeedbackRequest(BaseModel):
    feedback_type: str
    message: str


def generate_flashcards_fallback(text: str):
    sentences = [s.strip() for s in re.split(r"[.!?]\s+", text) if len(s.strip()) > 20]
    cards = []
    for sentence in sentences:
        if len(cards) >= 10:
            break
        front = sentence[:80].rstrip(".?!") + "?"
        back = sentence[:200]
        cards.append({"front": front, "back": back})
    if not cards:
        cards = [{"front": "Key topic", "back": text[:200]}]
    return cards


def extract_text_from_pdf(content: bytes) -> str:
    try:
        pdf = PyPDF2.PdfReader(io.BytesIO(content))
        text = ""
        page_count = len(pdf.pages)
        print(f"[DEBUG] PDF has {page_count} pages")

        for i, page in enumerate(pdf.pages):
            page_text = page.extract_text() or ""
            text += page_text + "\n"
            if i < 3:  # Log first 3 pages for debugging
                print(f"[DEBUG] Page {i+1} extracted {len(page_text)} characters")

        # Check if extraction was successful
        if len(text.strip()) < 50:
            print(f"[WARNING] PDF text extraction yielded only {len(text.strip())} characters from {page_count} pages")
            raise HTTPException(
                status_code=400,
                detail="Unable to extract text from PDF. The file may be scanned, image-based, or password-protected. Please try a text-based PDF or use an image format instead."
            )

        print(f"[DEBUG] Total extracted text: {len(text)} characters")
        return text
    except HTTPException:
        raise  # Re-raise our custom exception
    except Exception as e:
        print(f"[ERROR] PDF extraction failed: {str(e)}")
        raise HTTPException(
            status_code=400,
            detail=f"Failed to process PDF file: {str(e)}"
        )


def extract_text_from_docx(content: bytes) -> str:
    try:
        file_stream = io.BytesIO(content)
        doc = Document(file_stream)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs)

        print(f"[DEBUG] DOCX has {len(doc.paragraphs)} paragraphs")
        print(f"[DEBUG] Extracted {len(text)} characters from DOCX")

        # Check if extraction was successful
        if len(text.strip()) < 50:
            raise HTTPException(
                status_code=400,
                detail="Unable to extract sufficient text from DOCX file. The file may be empty or corrupted."
            )

        return text
    except HTTPException:
        raise  # Re-raise our custom exception
    except Exception as e:
        print(f"[ERROR] DOCX extraction failed: {str(e)}")
        raise HTTPException(
            status_code=400,
            detail=f"Failed to process DOCX file: {str(e)}"
        )


def generate_summary_from_text(text: str) -> str:
    if not os.getenv("OPENAI_API_KEY"):
        raise HTTPException(status_code=500, detail="OPENAI_API_KEY not configured")

    # Handle empty or very short text
    if not text or len(text.strip()) < 50:
        raise HTTPException(status_code=400, detail="Document contains insufficient content to summarize")

    # For short documents (under 12,000 chars), use single-pass summarization
    if len(text) <= 12000:
        prompt = """You are a study assistant. Summarize the notes clearly and concisely.
Return 5-8 bullet points plus a short 1-2 sentence overview.
Focus on key concepts, definitions, and important facts.
"""
        try:
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": prompt},
                    {"role": "user", "content": text}
                ],
                temperature=0.3,
                max_tokens=900
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            _raise_if_openai_error(e)
            raise

    # For long documents, use chunked summarization with map-reduce approach
    print(f"[DEBUG] Large document detected ({len(text)} chars). Using chunked summarization.")

    # Split text into chunks of ~10,000 characters, breaking at paragraph boundaries
    chunk_size = 10000
    chunks = []
    start = 0

    while start < len(text):
        end = start + chunk_size

        # If this isn't the last chunk, try to break at a paragraph or sentence boundary
        if end < len(text):
            # Look for paragraph break (double newline)
            paragraph_break = text.rfind("\n\n", start, end)
            if paragraph_break > start + chunk_size // 2:  # Only use if it's in the latter half of the chunk
                end = paragraph_break
            else:
                # Fall back to sentence break (period followed by space or newline)
                sentence_break = text.rfind(". ", start, end)
                if sentence_break > start + chunk_size // 2:
                    end = sentence_break + 1

        chunks.append(text[start:end])
        start = end

    print(f"[DEBUG] Split into {len(chunks)} chunks for summarization")

    # Generate summaries for each chunk
    chunk_summaries = []
    for i, chunk in enumerate(chunks):
        chunk_prompt = f"""You are a study assistant. This is part {i+1} of {len(chunks)} from a larger document.
Summarize this section clearly and concisely, focusing on key concepts, definitions, and important facts.
Return 3-5 bullet points covering the main ideas in this section.
"""
        try:
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": chunk_prompt},
                    {"role": "user", "content": chunk}
                ],
                temperature=0.3,
                max_tokens=500
            )
            chunk_summaries.append(response.choices[0].message.content.strip())
        except Exception as e:
            _raise_if_openai_error(e)
            raise

    # Combine chunk summaries into final comprehensive summary
    combined_text = "\n\n".join([f"Section {i+1}:\n{summary}" for i, summary in enumerate(chunk_summaries)])

    final_prompt = """You are a study assistant. Below are summaries of different sections from a larger document.
Create a unified, comprehensive summary that:
1. Provides a 2-3 sentence overview of the entire document
2. Lists 6-10 bullet points covering the most important concepts across all sections
3. Maintains logical flow and removes redundancy

Focus on the key takeaways a student needs to know for studying.
"""

    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": final_prompt},
                {"role": "user", "content": combined_text}
            ],
            temperature=0.3,
            max_tokens=1200
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        _raise_if_openai_error(e)
        raise


def generate_summary_from_image(content: bytes, filename: str) -> str:
    if not os.getenv("OPENAI_API_KEY"):
        raise HTTPException(status_code=500, detail="OPENAI_API_KEY not configured")

    ext = filename.lower().split(".")[-1]
    if ext == "jpg":
        ext = "jpeg"
    data_url = f"data:image/{ext};base64,{base64.b64encode(content).decode('utf-8')}"

    prompt = "Summarize the handwritten notes or study material in this image. Return 5-8 bullet points plus a short 1-2 sentence overview."

    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {"type": "image_url", "image_url": {"url": data_url}}
                    ],
                }
            ],
            temperature=0.3,
            max_tokens=900
        )
    except Exception as e:
        _raise_if_openai_error(e)
        raise

    return response.choices[0].message.content.strip()


# Create tables / migrations
# Ensure schema exists before running column updates
Base.metadata.create_all(bind=engine)
ensure_user_columns()
Base.metadata.create_all(bind=engine)


def ensure_indexes():
    """Create minimal indexes/constraints for user isolation and lookup performance."""
    index_statements = [
        # user-owned tables
        "CREATE INDEX IF NOT EXISTS idx_courses_user_id ON courses(user_id)",
        "CREATE INDEX IF NOT EXISTS idx_deadlines_user_id ON deadlines(user_id)",
        "CREATE INDEX IF NOT EXISTS idx_flashcard_sets_user_id ON flashcard_sets(user_id)",
        "CREATE INDEX IF NOT EXISTS idx_flashcards_user_id ON flashcards(user_id)",
        "CREATE INDEX IF NOT EXISTS idx_summaries_user_id ON summaries(user_id)",
        "CREATE INDEX IF NOT EXISTS idx_quizzes_user_id ON quizzes(user_id)",
        "CREATE INDEX IF NOT EXISTS idx_quiz_questions_user_id ON quiz_questions(user_id)",
        "CREATE INDEX IF NOT EXISTS idx_calendar_entries_user_id ON calendar_entries(user_id)",
        # user email uniqueness
        "CREATE UNIQUE INDEX IF NOT EXISTS idx_users_email ON users(email)",
    ]
    with engine.begin() as conn:
        for stmt in index_statements:
            try:
                conn.execute(text(stmt))
            except Exception as e:
                print(f"[WARN] Failed to apply index: {stmt} ({e})")


ensure_indexes()


def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


def _get_current_user(
    request: Request,
    credentials: HTTPAuthorizationCredentials | None,
    db,
) -> User:
    if credentials is None:
        raise HTTPException(status_code=401, detail="Not authenticated")

    token = credentials.credentials
    claims = _verify_supabase_token(token)
    user_id = claims.get("sub")
    email = claims.get("email")
    if not user_id:
        raise HTTPException(status_code=401, detail="Invalid token")

    user = db.query(User).filter(User.id == user_id).first()
    if not user:
        user = User(id=user_id, email=email or f"{user_id}@supabase.local")
        db.add(user)
        db.commit()
        db.refresh(user)
    elif email and user.email != email:
        user.email = email
        db.commit()
        db.refresh(user)

    return user


def get_current_user(
    request: Request,
    credentials: HTTPAuthorizationCredentials = Depends(bearer_scheme),
    db=Depends(get_db),
) -> User:
    return _get_current_user(request, credentials, db)


async def validate_file_upload(file: UploadFile, allowed_extensions: list[str], max_size_mb: int = 10):
    """
    Validate uploaded file for security and size constraints.

    Args:
        file: The uploaded file
        allowed_extensions: List of allowed file extensions (e.g., ['.pdf', '.docx'])
        max_size_mb: Maximum file size in megabytes

    Raises:
        HTTPException: If validation fails
    """
    # Validate file is provided
    if not file or not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")

    # Validate file extension
    file_ext = os.path.splitext(file.filename)[1].lower()
    if file_ext not in allowed_extensions:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid file type. Allowed types: {', '.join(allowed_extensions)}"
        )

    # Read file content to check size
    content = await file.read()
    file_size_mb = len(content) / (1024 * 1024)

    if file_size_mb > max_size_mb:
        raise HTTPException(
            status_code=400,
            detail=f"File too large. Maximum size: {max_size_mb}MB, uploaded: {file_size_mb:.2f}MB"
        )

    # Reset file pointer for subsequent reads
    await file.seek(0)

    return content


def parse_json_response(result: str):
    """Parse JSON from OpenAI response, handling markdown code blocks."""
    if result.startswith("```"):
        result = result.split("```")[1]
        if result.startswith("json"):
            result = result[4:]
    return json.loads(result.strip())


def extract_course_metadata(text: str):
    """PASS 1: Extract course metadata from syllabus."""
    print("[DEBUG] PASS 1: Extracting course metadata...")

    system_prompt = """You are a syllabus parser. Extract the course metadata AND detailed course information from this syllabus.

Return ONLY valid JSON with no additional text:
{
    "course_name": "Course code and name (e.g., 'FINC 313 - Corporate Finance')",
    "semester": "Semester name (e.g., 'Spring 2026')",
    "start_date": "First day of class in YYYY-MM-DD format (infer from semester if not explicit)",
    "end_date": "Last day of class/finals in YYYY-MM-DD format (infer from semester if not explicit)",
    "holidays": ["List of break dates or holidays mentioned"],
    "instructor": "Professor name if mentioned",
    "course_info": {
        "instructor": {
            "name": "Full name with title (e.g., 'Dr. John Smith')",
            "email": "Email address or null",
            "office": "Office location or null",
            "office_hours": "Office hours days/times or null",
            "phone": "Phone number or null"
        },
        "logistics": {
            "meeting_times": "Class meeting days/times (e.g., 'Mon/Wed 2:00-3:15 PM') or null",
            "location": "Classroom/building location or null",
            "attendance_policy": "Attendance policy summary or null",
            "late_work_policy": "Late work policy summary or null"
        },
        "grade_breakdown": [
            {"component": "Component name (e.g., 'Exams')", "weight": "Weight (e.g., '40%')"}
        ],
        "policies": {
            "participation": "Participation requirements or null",
            "extra_credit": "Extra credit opportunities or null",
            "academic_integrity": "Academic integrity policy summary or null",
            "prerequisites": "Prerequisite courses or null"
        },
        "materials": {
            "required_textbooks": ["List of required textbook titles"],
            "recommended_readings": ["List of recommended reading titles"],
            "course_portal": "Course website/portal URL or null",
            "ta_info": "TA name and contact info or null"
        }
    }
}

For semester dates, use these typical academic calendars:
- Spring semester: mid-January to mid-May
- Fall semester: late August to mid-December
- Summer: May to August

If the year is 2026 and semester is Spring, start_date would be around 2026-01-12 and end_date around 2026-05-08.

IMPORTANT: For any field where information is not found in the syllabus, use null (for strings) or empty arrays (for lists). Extract as much detail as possible."""

    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Extract metadata from this syllabus:\n\n{text[:15000]}"}
            ],
            temperature=0.1,
            max_tokens=2000
        )

        result = response.choices[0].message.content
        print(f"[DEBUG] Metadata response: {result}")
        metadata = parse_json_response(result)
        print(f"[DEBUG] Parsed metadata: {metadata}")
        return metadata

    except Exception as e:
        _raise_if_openai_error(e)  # don't silently swallow auth / rate-limit errors
        print(f"[ERROR] Metadata extraction failed: {e}")
        return {
            "course_name": "Unknown Course",
            "semester": "Unknown",
            "start_date": None,
            "end_date": None,
            "holidays": [],
            "instructor": None,
            "course_info": None
        }


def extract_deadlines_with_context(text: str, metadata: dict):
    """PASS 2: Extract deadlines using course metadata for context."""
    print("[DEBUG] PASS 2: Extracting deadlines with context...")

    max_chars = 25000
    if len(text) > max_chars:
        text = text[:max_chars]
        print(f"[DEBUG] Truncated text to {max_chars} characters")

    start_date = metadata.get("start_date", "2026-01-12")
    end_date = metadata.get("end_date", "2026-05-08")
    semester = metadata.get("semester", "Spring 2026")

    system_prompt = f"""You are parsing a college course syllabus. The course runs from {start_date} to {end_date} ({semester}).

EXTRACT THESE TYPES OF DEADLINES (be thorough - extract ALL you find):

1. EXAMS & TESTS - Midterms, finals, tests with SPECIFIC DATES (look for "Test 1", "Exam 2", "Final Exam", "Final Test")
2. QUIZZES - Pop quizzes, scheduled quizzes with dates (look for "Quiz 1", "Quiz 2", etc.)
3. MAJOR ASSIGNMENTS - Papers, projects, case studies with DUE DATES
4. HOMEWORK - HW assignments with due dates (look for "HW 1 DUE", "Homework due", etc.)
5. PRESENTATIONS - Pitches, presentations with specific dates (look for "Pitch 1", "Mini Pitch", "Presentation")
6. RECURRING ASSESSMENTS - Weekly quizzes, homework, reading checks (mark as recurring)
7. IMPORTANT ADMIN DATES - Add/drop deadline, makeup day, last day of class, final assignment deadline

IMPORTANT - LOOK FOR THESE SECTIONS:
- "IMPORTANT DAYS" or "IMPORTANT DATES" sections - these list key deadlines
- Schedule tables with columns like "Session", "Day", "Topic", "Notes" - extract quiz/test/HW dates from these
- "TENTATIVE SCHEDULE" sections with dated items
- Lines like "Session 11 Quiz 1 01/30/2026" or "Quiz 1: 01/30/2026"
- Homework assignments with dates like "HW 1 IS GIVEN" followed by "HW 1 DUE"
- Final exam schedules with dates and times
- "All assignments due by [date]" - this is a major deadline

DATE FORMATS TO RECOGNIZE:
- MM/DD/YYYY (e.g., 01/30/2026)
- Month DD, YYYY (e.g., January 30, 2026)
- DD-Mon (e.g., 30-Jan)
- Dates in tables (look for patterns like "7-Jan", "9-Jan", etc.)

DO NOT EXTRACT:
- Regular class meeting times (unless it's also an exam/quiz day)
- Office hours
- Reading assignments without assessments
- Topic lists without deliverables
- Spring break, holidays (unless they're makeup days)

HANDLING RECURRING ITEMS:
If you see patterns like "Quiz every Monday", "Weekly homework due Fridays":
- Create ONE entry with recurring=true
- Set frequency="weekly" and day_of_week to the specific day
- Use the FIRST occurrence date

Return ONLY a valid JSON array. Each item must have:
{{
    "date": "YYYY-MM-DD (first occurrence for recurring, or specific date)",
    "type": "Exam|Assignment|Project|Quiz|Homework|Presentation|Admin",
    "title": "Descriptive name (e.g., 'Test 1', 'Quiz 1', 'Mini Pitch 1', 'HW 1 Due')",
    "context": "Brief description from syllabus",
    "time": "Due time if mentioned (e.g., '11:59pm', '7:00 AM'), or null",
    "recurring": true/false (true if it repeats weekly),
    "frequency": "weekly" or null,
    "day_of_week": "Monday|Tuesday|...|Sunday" or null
}}

Examples of GOOD entries:
- {{"date": "2026-01-30", "type": "Quiz", "title": "Quiz 1", "context": "Covers chapters 1-2-3", "time": null, "recurring": false}}
- {{"date": "2026-02-09", "type": "Exam", "title": "Test 1", "context": "Covers chapters 1-2-3-4", "time": null, "recurring": false}}
- {{"date": "2026-01-29", "type": "Presentation", "title": "Mini Pitch 1", "context": "Customer discovery results presentation", "time": null, "recurring": false}}
- {{"date": "2026-04-22", "type": "Admin", "title": "Final Assignment Due", "context": "All assignments due, no extensions", "time": "7:00 AM", "recurring": false}}
- {{"date": "2026-04-24", "type": "Exam", "title": "Final Exam", "context": "Final test for section 6", "time": "8:00 AM", "recurring": false}}
- {{"date": "2026-01-16", "type": "Homework", "title": "HW 1 Due", "context": "First homework assignment due", "time": "11:59pm", "recurring": false}}

Return [] if no deadlines found."""

    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Extract ALL deadlines, quizzes, tests, exams, homework due dates, presentations, and important dates from this syllabus. Pay special attention to 'IMPORTANT DAYS' sections, schedule tables, and any dates with Quiz/Test/Exam/HW/Pitch/Presentation labels:\n\n{text}"}
            ],
            temperature=0.1,
            max_tokens=6000
        )

        result = response.choices[0].message.content
        print(f"[DEBUG] Deadlines response (first 1000 chars): {result[:1000]}...")

        deadlines = parse_json_response(result)
        print(f"[DEBUG] Parsed {len(deadlines)} deadlines before validation")

        # Validate and filter deadlines
        validated = validate_deadlines(deadlines)
        print(f"[DEBUG] {len(validated)} deadlines after validation")

        # Expand recurring deadlines into individual instances
        expanded = expand_recurring_deadlines(validated, start_date, end_date)
        print(f"[DEBUG] {len(expanded)} deadlines after recurring expansion")

        return expanded

    except json.JSONDecodeError as e:
        print(f"[ERROR] Failed to parse deadlines JSON: {e}")
        return []
    except Exception as e:
        _raise_if_openai_error(e)
        print(f"[ERROR] Deadline extraction failed: {e}")
        raise


def expand_recurring_deadlines(deadlines: list, start_date_str: str, end_date_str: str) -> list:
    """Expand recurring deadlines into individual instances."""
    DAY_MAP = {
        "monday": 0, "tuesday": 1, "wednesday": 2, "thursday": 3,
        "friday": 4, "saturday": 5, "sunday": 6
    }

    try:
        start = datetime.strptime(start_date_str, "%Y-%m-%d").date()
        end = datetime.strptime(end_date_str, "%Y-%m-%d").date()
    except (ValueError, TypeError):
        print("[DEBUG] Invalid date range for recurring expansion, using defaults")
        start = date(2026, 1, 12)
        end = date(2026, 5, 8)

    expanded = []

    for d in deadlines:
        recurring = d.get("recurring", False)
        frequency = (d.get("frequency") or "").lower()
        day_of_week = (d.get("day_of_week") or "").lower()

        if recurring and frequency == "weekly" and day_of_week in DAY_MAP:
            # Expand weekly recurring deadline
            target_day = DAY_MAP[day_of_week]
            current = start
            instance_num = 1

            # Find first occurrence of the target day
            while current.weekday() != target_day:
                current += timedelta(days=1)

            # Generate all instances (skip first week usually just intro)
            current += timedelta(days=7)

            while current <= end:
                instance = {
                    "date": current.strftime("%Y-%m-%d"),
                    "time": d.get("time"),
                    "type": d.get("type", "Quiz"),
                    "title": f"{d.get('title', 'Weekly Item')} #{instance_num}",
                    "context": d.get("context") or d.get("description"),
                    "recurring": True,
                    "frequency": "weekly",
                    "day_of_week": d.get("day_of_week")
                }
                expanded.append(instance)
                instance_num += 1
                current += timedelta(days=7)

                # Limit to 15 instances max
                if instance_num > 15:
                    break

            print(f"[DEBUG] Expanded recurring '{d.get('title')}' into {instance_num - 1} instances")
        else:
            # Not recurring, keep as-is
            expanded.append(d)

    return expanded


def validate_deadlines(deadlines: list) -> list:
    """Filter out low-quality or duplicate deadlines."""
    validated = []
    seen_titles = set()

    # Generic titles to reject (but allow if marked as recurring with specific day)
    generic_patterns = [
        "class meeting",
        "office hours",
        "lecture",
    ]

    # Week patterns to reject
    week_patterns = ["week 1", "week 2", "week 3", "week 4", "week 5"]

    for d in deadlines:
        title = (d.get("title") or "").lower().strip()
        date_str = d.get("date") or ""

        # Skip if no date
        if not date_str or date_str == "null":
            print(f"[DEBUG] Skipping deadline with no date: {title}")
            continue

        # Skip class meetings, office hours, lectures
        if any(pattern in title for pattern in generic_patterns):
            print(f"[DEBUG] Skipping generic deadline: {title}")
            continue

        # Skip plain "Week X" entries
        if any(pattern in title for pattern in week_patterns):
            print(f"[DEBUG] Skipping week entry: {title}")
            continue

        # Skip duplicates (same title + same date)
        title_key = f"{title[:30]}_{date_str}"
        if title_key in seen_titles:
            print(f"[DEBUG] Skipping duplicate: {title}")
            continue
        seen_titles.add(title_key)

        validated.append(d)

    return validated


@app.on_event("startup")
async def startup_event():
    """Initialize app: validate env vars, test DB, pre-fetch JWKS, log configuration."""
    logger.info("=" * 50)
    logger.info("[Startup] ClassMate Backend API starting...")
    logger.info("=" * 50)

    # Validate required environment variables
    required_vars = ["OPENAI_API_KEY", "SUPABASE_URL"]
    missing = [var for var in required_vars if not os.getenv(var)]
    if missing:
        logger.error(f"[Startup] FATAL: Missing required env vars: {missing}")
        raise RuntimeError(f"Missing required environment variables: {missing}")

    # Log CORS configuration
    logger.info(f"[Startup] CORS enabled for origins: {ALLOWED_ORIGINS}")

    # Test database connection
    try:
        with engine.connect() as conn:
            conn.execute(text("SELECT 1"))
        logger.info("[Startup] Database connection successful")
        logger.info(f"[Startup] Database URL: {_safe_db_url(DATABASE_URL)}")
    except Exception as e:
        logger.error(f"[Startup] ERROR: Database connection failed: {e}")

    # Pre-fetch JWKS
    logger.info("[Startup] Pre-fetching JWKS...")
    keys = _fetch_jwks()
    if keys:
        logger.info(f"[Startup] JWKS pre-fetch successful: {len(keys)} keys cached")
    else:
        logger.warning("[Startup] JWKS pre-fetch failed - auth may not work until keys can be fetched")
        logger.warning(f"[Startup] SUPABASE_URL = {SUPABASE_URL or 'NOT SET'}")
        logger.warning(f"[Startup] SUPABASE_JWKS_URL = {SUPABASE_JWKS_URL or 'NOT SET'}")

    logger.info("=" * 50)
    logger.info("[Startup] ClassMate Backend API ready!")
    logger.info("=" * 50)


@app.get("/")
def root():
    """Root endpoint - service status."""
    return {
        "service": "ClassMate Backend API",
        "status": "running",
        "version": "1.0.0"
    }


@app.get("/health")
def health():
    """Health check endpoint."""
    return {"status": "healthy"}


@app.get("/api/health")
def api_health():
    """API health check endpoint."""
    return {"status": "healthy"}


@app.get("/auth-status")
def auth_status():
    """Debug endpoint to check auth configuration status."""
    cached_keys = _jwks_cache.get("keys", {})
    fetched_at = _jwks_cache.get("fetched_at", 0)
    last_error = _jwks_cache.get("last_fetch_error")
    return {
        "supabase_url_configured": bool(SUPABASE_URL),
        "jwks_url_configured": bool(SUPABASE_JWKS_URL),
        "issuer_configured": bool(SUPABASE_ISSUER),
        "jwks_keys_cached": len(cached_keys),
        "jwks_cache_age_seconds": int(time.time() - float(fetched_at)) if fetched_at else None,
        "last_fetch_error": last_error,
    }


@app.post("/courses")
def create_course(payload: CreateCourseRequest, current_user: User = Depends(get_current_user)):
    """Create a course manually."""
    db = SessionLocal()
    try:
        user_id = current_user.id
        if not payload.name:
            raise HTTPException(status_code=400, detail="Course name is required")

        start_date = None
        end_date = None
        if payload.start_date:
            try:
                start_date = date.fromisoformat(payload.start_date)
            except ValueError:
                raise HTTPException(status_code=400, detail="Invalid start_date format")
        if payload.end_date:
            try:
                end_date = date.fromisoformat(payload.end_date)
            except ValueError:
                raise HTTPException(status_code=400, detail="Invalid end_date format")

        course = Course(
            user_id=user_id,
            name=payload.name.strip(),
            code=payload.code.strip() if payload.code else None,
            semester=payload.semester.strip() if payload.semester else "Term TBD",
            start_date=start_date,
            end_date=end_date,
        )
        db.add(course)
        db.commit()
        db.refresh(course)

        return {
            "id": course.id,
            "name": course.name,
            "code": course.code,
            "semester": course.semester,
            "start_date": str(course.start_date) if course.start_date else None,
            "end_date": str(course.end_date) if course.end_date else None,
            "course_info": course.course_info,
            "deadline_count": 0,
            "flashcard_set_count": 0,
            "created_at": course.created_at.isoformat(),
        }
    finally:
        db.close()


def _resolve_profile(db, user_id: str, email: str | None = None) -> UserProfile | None:
    """Look up a UserProfile by user_id first, then fall back to email.

    If a profile is found via email but its user_id doesn't match the current
    JWT sub, we update it in place so that all future lookups (and every other
    table that references user_id) can be migrated in one go later.  This is a
    one-time self-heal per profile — after the first successful request the
    user_id column will already be correct.
    """
    profile = db.query(UserProfile).filter(UserProfile.user_id == user_id).first()
    if profile:
        return profile
    # Fallback: match on email and heal the stale user_id
    if email:
        profile = db.query(UserProfile).filter(UserProfile.email == email).first()
        if profile:
            profile.user_id = user_id
            db.commit()
            db.refresh(profile)
    return profile


@app.get("/me")
def get_me(current_user: User = Depends(get_current_user)):
    db = SessionLocal()
    try:
        profile = _resolve_profile(db, current_user.id, current_user.email)
        if not profile:
            return {"user_id": current_user.id, "profile": None}
        return {
            "user_id": current_user.id,
            "profile": {
                "email": profile.email,
                "full_name": profile.full_name,
                "school_name": profile.school_name,
                "school_type": profile.school_type,
                "academic_year": profile.academic_year,
                "major": profile.major,
                "profile_picture": profile.profile_picture,
            }
        }
    finally:
        db.close()


@app.post("/me/profile")
def upsert_profile(payload: UserProfileRequest, current_user: User = Depends(get_current_user)):
    db = SessionLocal()
    try:
        profile = _resolve_profile(db, current_user.id, current_user.email)
        if not profile:
            # Create the profile row so new users are never stuck
            profile = UserProfile(
                user_id=current_user.id,
                email=current_user.email or (payload.email or ""),
                full_name=payload.full_name,
                school_name=payload.school_name,
                school_type=payload.school_type,
                academic_year=payload.academic_year,
                major=payload.major,
            )
            db.add(profile)
            db.commit()
            db.refresh(profile)
        else:
            profile.full_name = payload.full_name or profile.full_name
            profile.school_name = payload.school_name or profile.school_name
            profile.school_type = payload.school_type or profile.school_type
            profile.academic_year = payload.academic_year or profile.academic_year
            profile.major = payload.major or profile.major
            db.commit()

        return {
            "message": "Profile saved",
            "profile": {
                "email": profile.email,
                "full_name": profile.full_name,
                "school_name": profile.school_name,
                "school_type": profile.school_type,
                "academic_year": profile.academic_year,
                "major": profile.major,
                "profile_picture": profile.profile_picture,
            }
        }
    finally:
        db.close()


@app.post("/me/profile-picture")
def upload_profile_picture(payload: ProfilePictureRequest, current_user: User = Depends(get_current_user)):
    """Upload a profile picture as base64."""
    db = SessionLocal()
    try:
        profile = _resolve_profile(db, current_user.id, current_user.email)
        if not profile:
            # Auto-create a minimal profile so the picture can be saved immediately
            profile = UserProfile(
                user_id=current_user.id,
                email=current_user.email or "",
            )
            db.add(profile)
            db.commit()
            db.refresh(profile)

        if len(payload.image_data) > 500_000:
            raise HTTPException(status_code=400, detail="Image too large (max ~375KB)")

        profile.profile_picture = payload.image_data
        db.commit()
        return {"message": "Profile picture updated", "profile_picture": profile.profile_picture}
    finally:
        db.close()


@app.get("/courses")
def list_courses(current_user: User = Depends(get_current_user)):
    """List all courses."""
    db = SessionLocal()
    try:
        user_id = current_user.id
        courses = db.query(Course).filter(Course.user_id == user_id).order_by(Course.created_at.desc()).all()
        return [
            {
                "id": c.id,
                "name": c.name,
                "code": c.code,
                "semester": c.semester,
                "start_date": str(c.start_date) if c.start_date else None,
                "end_date": str(c.end_date) if c.end_date else None,
                "course_info": c.course_info,
                "deadline_count": len(c.deadlines),
                "flashcard_set_count": len(c.flashcard_sets),
                "created_at": c.created_at.isoformat()
            }
            for c in courses
        ]
    finally:
        db.close()


@app.get("/courses/{course_id}")
def get_course(course_id: str, current_user: User = Depends(get_current_user)):
    """Get a course with its deadlines and flashcard sets."""
    db = SessionLocal()
    try:
        user_id = current_user.id
        course = db.query(Course).filter(Course.id == course_id, Course.user_id == user_id).first()
        if not course:
            raise HTTPException(status_code=404, detail="Course not found")

        # Get calendar entries for this course's deadlines
        deadline_ids = [d.id for d in course.deadlines]
        calendar_entries = db.query(CalendarEntry).filter(
            CalendarEntry.deadline_id.in_(deadline_ids),
            CalendarEntry.user_id == user_id
        ).all()
        saved_deadline_ids = {e.deadline_id for e in calendar_entries}

        return {
            "id": course.id,
            "name": course.name,
            "code": course.code,
            "semester": course.semester,
            "start_date": str(course.start_date) if course.start_date else None,
            "end_date": str(course.end_date) if course.end_date else None,
            "course_info": course.course_info,
            "deadlines": [
                {
                    "id": d.id,
                    "date": d.date,
                    "time": d.time,
                    "type": d.type,
                    "title": d.title,
                    "description": d.description,
                    "recurring": d.recurring,
                    "frequency": d.frequency,
                    "day_of_week": d.day_of_week,
                    "completed": d.completed,
                    "saved_to_calendar": d.id in saved_deadline_ids
                }
                for d in sorted(course.deadlines, key=lambda x: x.date or "9999")
            ],
            "flashcard_sets": [
                {
                    "id": fs.id,
                    "name": fs.name,
                    "card_count": len(fs.flashcards),
                    "created_at": fs.created_at.isoformat()
                }
                for fs in course.flashcard_sets
            ],
            "summaries": [
                {
                    "id": s.id,
                    "title": s.title,
                    "content": s.content,
                    "created_at": s.created_at.isoformat()
                }
                for s in sorted(course.summaries, key=lambda x: x.created_at, reverse=True)
            ],
            "quizzes": [
                {
                    "id": q.id,
                    "name": q.name,
                    "question_count": len(q.questions),
                    "created_at": q.created_at.isoformat()
                }
                for q in course.quizzes
            ]
        }
    finally:
        db.close()


@app.get("/courses/{course_id}/deadlines")
def list_course_deadlines(course_id: str, current_user: User = Depends(get_current_user)):
    """List deadlines for a specific course."""
    db = SessionLocal()
    try:
        user_id = current_user.id
        deadlines = db.query(Deadline).filter(Deadline.course_id == course_id, Deadline.user_id == user_id).all()
        return [
            {
                "id": d.id,
                "course_id": d.course_id,
                "date": d.date,
                "time": d.time,
                "type": d.type,
                "title": d.title,
                "description": d.description,
                "recurring": d.recurring,
                "frequency": d.frequency,
                "day_of_week": d.day_of_week,
                "completed": d.completed
            }
            for d in sorted(deadlines, key=lambda x: x.date or "9999")
        ]
    finally:
        db.close()


@app.get("/deadlines")
def list_all_deadlines(
    from_date: str | None = Query(default=None, alias="from"),
    to_date: str | None = Query(default=None, alias="to"),
    course_id: str | None = None,
    current_user: User = Depends(get_current_user),
):
    """Get all deadlines from all courses (for calendar view)."""
    db = SessionLocal()
    try:
        user_id = current_user.id
        print(f"[DEBUG] /deadlines request from={from_date} to={to_date} course_id={course_id}")
        query = db.query(Deadline).join(Course).filter(Deadline.user_id == user_id)
        if course_id:
            query = query.filter(Deadline.course_id == course_id)
        if from_date:
            query = query.filter(Deadline.date >= from_date)
        if to_date:
            query = query.filter(Deadline.date <= to_date)
        deadlines = query.all()
        return [
            {
                "id": d.id,
                "course_id": d.course_id,
                "course_name": d.course.name,
                "course_code": d.course.code,
                "date": d.date,
                "time": d.time,
                "type": d.type,
                "title": d.title,
                "description": d.description,
                "recurring": d.recurring,
                "frequency": d.frequency,
                "day_of_week": d.day_of_week,
                "completed": d.completed
            }
            for d in sorted(deadlines, key=lambda x: x.date or "9999")
        ]
    finally:
        db.close()


@app.post("/deadlines")
def create_deadline(payload: CreateDeadlineRequest, current_user: User = Depends(get_current_user)):
    """Create a custom deadline (user-generated, not from syllabus)."""
    db = SessionLocal()
    try:
        user_id = current_user.id
        if not payload.course_id:
            raise HTTPException(status_code=400, detail="course_id is required")

        course = db.query(Course).filter(Course.id == payload.course_id, Course.user_id == user_id).first()
        if not course:
            raise HTTPException(status_code=404, detail="Course not found")

        deadline = Deadline(
            user_id=user_id,
            course_id=payload.course_id,
            date=payload.date,
            time=payload.time,
            type=payload.type or "Deadline",
            title=payload.title,
            description=payload.description,
            recurring=False,
        )
        db.add(deadline)
        db.commit()
        db.refresh(deadline)

        return {
            "id": deadline.id,
            "course_id": deadline.course_id,
            "course_name": course.name if course else None,
            "course_code": course.code if course else None,
            "date": deadline.date,
            "time": deadline.time,
            "type": deadline.type,
            "title": deadline.title,
            "description": deadline.description,
            "completed": deadline.completed,
        }
    finally:
        db.close()


@app.delete("/deadlines/{deadline_id}")
def delete_deadline(deadline_id: str, current_user: User = Depends(get_current_user)):
    """Delete a deadline."""
    db = SessionLocal()
    try:
        user_id = current_user.id
        deadline = db.query(Deadline).filter(Deadline.id == deadline_id, Deadline.user_id == user_id).first()
        if not deadline:
            raise HTTPException(status_code=404, detail="Deadline not found")

        # Also remove calendar entry if exists
        calendar_entry = db.query(CalendarEntry).filter(
            CalendarEntry.deadline_id == deadline_id,
            CalendarEntry.user_id == user_id
        ).first()
        if calendar_entry:
            db.delete(calendar_entry)

        db.delete(deadline)
        db.commit()
        return {"message": "Deadline deleted"}
    finally:
        db.close()


@app.patch("/deadlines/{deadline_id}/complete")
def toggle_deadline_complete(deadline_id: str, current_user: User = Depends(get_current_user)):
    """Toggle a deadline's completed status."""
    db = SessionLocal()
    try:
        user_id = current_user.id
        deadline = db.query(Deadline).filter(Deadline.id == deadline_id, Deadline.user_id == user_id).first()
        if not deadline:
            raise HTTPException(status_code=404, detail="Deadline not found")

        deadline.completed = not deadline.completed
        db.commit()
        return {"id": deadline.id, "completed": deadline.completed}
    finally:
        db.close()


@app.patch("/deadlines/{deadline_id}")
def update_deadline(deadline_id: str, payload: UpdateDeadlineRequest, current_user: User = Depends(get_current_user)):
    """Update deadline fields (currently supports date only)."""
    db = SessionLocal()
    try:
        user_id = current_user.id
        deadline = db.query(Deadline).filter(Deadline.id == deadline_id, Deadline.user_id == user_id).first()
        if not deadline:
            raise HTTPException(status_code=404, detail="Deadline not found")

        if payload.date:
            deadline.date = payload.date

        db.commit()
        return {"id": deadline.id, "date": deadline.date}
    finally:
        db.close()

@app.post("/deadlines/{deadline_id}/save-to-calendar")
def save_deadline_to_calendar(deadline_id: str, current_user: User = Depends(get_current_user)):
    """Save a deadline to the user's calendar."""
    db = SessionLocal()
    try:
        user_id = current_user.id
        deadline = db.query(Deadline).filter(Deadline.id == deadline_id, Deadline.user_id == user_id).first()
        if not deadline:
            raise HTTPException(status_code=404, detail="Deadline not found")

        # Check if already saved
        existing = db.query(CalendarEntry).filter(
            CalendarEntry.deadline_id == deadline_id,
            CalendarEntry.user_id == user_id
        ).first()
        if existing:
            return {"id": existing.id, "deadline_id": deadline_id, "already_saved": True}

        # Create calendar entry
        entry = CalendarEntry(deadline_id=deadline_id, user_id=user_id)
        db.add(entry)
        db.commit()
        db.refresh(entry)

        return {
            "id": entry.id,
            "deadline_id": deadline_id,
            "created_at": entry.created_at.isoformat(),
            "already_saved": False
        }
    finally:
        db.close()


@app.delete("/deadlines/{deadline_id}/save-to-calendar")
def remove_deadline_from_calendar(deadline_id: str, current_user: User = Depends(get_current_user)):
    """Remove a deadline from the user's calendar."""
    db = SessionLocal()
    try:
        user_id = current_user.id
        entry = db.query(CalendarEntry).filter(
            CalendarEntry.deadline_id == deadline_id,
            CalendarEntry.user_id == user_id
        ).first()
        if not entry:
            raise HTTPException(status_code=404, detail="Calendar entry not found")

        db.delete(entry)
        db.commit()
        return {"message": "Removed from calendar"}
    finally:
        db.close()


@app.get("/calendar-entries")
def list_calendar_entries(current_user: User = Depends(get_current_user)):
    """Get all deadlines that have been saved to calendar."""
    db = SessionLocal()
    try:
        user_id = current_user.id
        entries = db.query(CalendarEntry).filter(CalendarEntry.user_id == user_id).all()
        result = []
        for entry in entries:
            deadline = entry.deadline
            if deadline:
                result.append({
                    "id": entry.id,
                    "deadline_id": deadline.id,
                    "course_id": deadline.course_id,
                    "course_name": deadline.course.name if deadline.course else None,
                    "course_code": deadline.course.code if deadline.course else None,
                    "date": deadline.date,
                    "time": deadline.time,
                    "type": deadline.type,
                    "title": deadline.title,
                    "description": deadline.description,
                    "completed": deadline.completed,
                    "saved_at": entry.created_at.isoformat()
                })
        return result
    finally:
        db.close()


@app.delete("/courses/{course_id}")
def delete_course(course_id: str, current_user: User = Depends(get_current_user)):
    """Delete a course and all its deadlines."""
    db = SessionLocal()
    try:
        user_id = current_user.id
        course = db.query(Course).filter(Course.id == course_id, Course.user_id == user_id).first()
        if not course:
            raise HTTPException(status_code=404, detail="Course not found")

        deadline_ids = [d.id for d in course.deadlines]
        if deadline_ids:
            db.query(CalendarEntry).filter(
                CalendarEntry.deadline_id.in_(deadline_ids),
                CalendarEntry.user_id == user_id
            ).delete(synchronize_session=False)

        db.delete(course)
        db.commit()
        return {"message": "Course deleted"}
    finally:
        db.close()


@app.patch("/courses/{course_id}")
def update_course(course_id: str, payload: UpdateCourseRequest, current_user: User = Depends(get_current_user)):
    """Update course fields."""
    db = SessionLocal()
    try:
        user_id = current_user.id
        course = db.query(Course).filter(Course.id == course_id, Course.user_id == user_id).first()
        if not course:
            raise HTTPException(status_code=404, detail="Course not found")

        if payload.name is not None:
            course.name = payload.name.strip()
        if payload.code is not None:
            course.code = payload.code.strip() if payload.code else None
        if payload.semester is not None:
            course.semester = payload.semester.strip() if payload.semester else None
        if payload.start_date:
            course.start_date = date.fromisoformat(payload.start_date)
        if payload.end_date:
            course.end_date = date.fromisoformat(payload.end_date)
        if payload.course_info is not None:
            course.course_info = payload.course_info

        db.commit()
        db.refresh(course)
        return {
            "id": course.id,
            "name": course.name,
            "code": course.code,
            "semester": course.semester,
            "start_date": str(course.start_date) if course.start_date else None,
            "end_date": str(course.end_date) if course.end_date else None,
            "course_info": course.course_info,
        }
    finally:
        db.close()


@app.post("/courses/{course_id}/syllabus")
@limiter.limit("5/minute")
async def upload_course_syllabus(request: Request, course_id: str, file: UploadFile = File(...), current_user: User = Depends(get_current_user)):
    """Upload a syllabus PDF or Word doc and attach extracted deadlines to an existing course."""
    logger.info(f"[DEBUG] /courses/{course_id}/syllabus request received")

    # Validate file upload — PDF and DOCX both supported
    content = await validate_file_upload(file, allowed_extensions=['.pdf', '.docx'], max_size_mb=10)
    filename = (file.filename or "").lower()

    if filename.endswith(".docx"):
        text = extract_text_from_docx(content)
    else:
        text = extract_text_from_pdf(content)

    print(f"[DEBUG] Total text extracted: {len(text)} characters")

    if len(text.strip()) < 50:
        raise HTTPException(status_code=400, detail="Could not extract text from the uploaded file. Make sure it contains readable text.")

    metadata = extract_course_metadata(text)
    deadlines_data = extract_deadlines_with_context(text, metadata)
    print(f"[DEBUG] Deadline extraction returned {len(deadlines_data)} items")

    db = SessionLocal()
    try:
        user_id = current_user.id
        course = db.query(Course).filter(Course.id == course_id, Course.user_id == user_id).first()
        if not course:
            raise HTTPException(status_code=404, detail="Course not found")

        # Update course_info from parsed metadata
        course_info = metadata.get("course_info")
        if course_info:
            course.course_info = course_info

        for d in deadlines_data:
            deadline = Deadline(
                user_id=user_id,
                course_id=course_id,
                date=d.get("date"),
                time=d.get("time"),
                type=d.get("type"),
                title=d.get("title"),
                description=d.get("context") or d.get("description"),
                recurring=d.get("recurring", False),
                frequency=d.get("frequency"),
                day_of_week=d.get("day_of_week")
            )
            db.add(deadline)

        db.commit()
        print(f"[DEBUG] Inserted {len(deadlines_data)} deadlines for course {course_id}")

        return {
            "course_id": course_id,
            "deadlines_created": len(deadlines_data),
            "deadlines": deadlines_data,
            "course_info": course.course_info
        }
    finally:
        db.close()


@app.delete("/courses/{course_id}/syllabus")
def delete_course_syllabus(course_id: str, current_user: User = Depends(get_current_user)):
    """Remove all AI-extracted deadlines and course_info for a course so the syllabus can be re-uploaded.

    Only deadlines that were created by syllabus extraction (i.e. all deadlines on the course)
    are removed.  Flashcard sets, summaries, and quizzes are NOT touched — those come from
    separate study-material uploads and have their own delete endpoints.
    """
    db = SessionLocal()
    try:
        user_id = current_user.id
        course = db.query(Course).filter(Course.id == course_id, Course.user_id == user_id).first()
        if not course:
            raise HTTPException(status_code=404, detail="Course not found")

        # Remove calendar entries that reference this course's deadlines first
        deadline_ids = [d.id for d in course.deadlines]
        if deadline_ids:
            db.query(CalendarEntry).filter(
                CalendarEntry.deadline_id.in_(deadline_ids),
                CalendarEntry.user_id == user_id
            ).delete(synchronize_session=False)

        # Delete all deadlines
        db.query(Deadline).filter(Deadline.course_id == course_id).delete()

        # Clear extracted course_info so the sidebar shows the syllabus upload zone again
        course.course_info = None
        db.commit()

        return {"message": "Syllabus and extracted deadlines removed", "course_id": course_id}
    finally:
        db.close()


@app.post("/courses/{course_id}/summaries")
@limiter.limit("10/minute")
async def generate_summary(request: Request, course_id: str, file: UploadFile = File(...), current_user: User = Depends(get_current_user)):
    """Upload notes and generate a study summary attached to a course."""
    logger.info(f"[DEBUG] /courses/{course_id}/summaries request received")
    db = SessionLocal()
    try:
        user_id = current_user.id
        course = db.query(Course).filter(Course.id == course_id, Course.user_id == user_id).first()
        if not course:
            raise HTTPException(status_code=404, detail="Course not found")

        # Validate file upload - allow PDF, TXT, and DOCX files
        content = await validate_file_upload(file, allowed_extensions=['.pdf', '.txt', '.docx'], max_size_mb=10)
        filename = file.filename.lower()

        if filename.endswith(".pdf"):
            text = extract_text_from_pdf(content)
            print(f"[DEBUG] Summary PDF text length: {len(text)}")
            summary_text = generate_summary_from_text(text)
        elif filename.endswith(".txt"):
            text = content.decode("utf-8", errors="ignore")
            print(f"[DEBUG] Summary TXT length: {len(text)}")
            summary_text = generate_summary_from_text(text)
        elif filename.endswith(".docx"):
            text = extract_text_from_docx(content)
            print(f"[DEBUG] Summary DOCX length: {len(text)}")
            summary_text = generate_summary_from_text(text)
        elif filename.endswith((".png", ".jpg", ".jpeg")):
            print(f"[DEBUG] Summary image size: {len(content)} bytes")
            summary_text = generate_summary_from_image(content, filename)
        else:
            raise HTTPException(status_code=400, detail="Supported formats: PDF, DOCX, TXT, PNG, JPG")

        summary = Summary(
            user_id=user_id,
            course_id=course_id,
            title=file.filename.rsplit(".", 1)[0],
            content=summary_text
        )
        db.add(summary)
        db.commit()
        db.refresh(summary)

        print(f"[DEBUG] Summary created {summary.id}")
        return {
            "id": summary.id,
            "course_id": course_id,
            "title": summary.title,
            "content": summary.content,
            "created_at": summary.created_at.isoformat()
        }
    finally:
        db.close()


@app.delete("/summaries/{summary_id}")
def delete_summary(summary_id: str, current_user: User = Depends(get_current_user)):
    """Delete a summary."""
    db = SessionLocal()
    try:
        user_id = current_user.id
        summary = db.query(Summary).filter(Summary.id == summary_id, Summary.user_id == user_id).first()
        if not summary:
            raise HTTPException(status_code=404, detail="Summary not found")
        db.delete(summary)
        db.commit()
        return {"message": "Summary deleted"}
    finally:
        db.close()


@app.get("/summaries/{summary_id}")
def get_summary(summary_id: str, current_user: User = Depends(get_current_user)):
    """Get a summary with course info."""
    db = SessionLocal()
    try:
        user_id = current_user.id
        summary = db.query(Summary).filter(Summary.id == summary_id, Summary.user_id == user_id).first()
        if not summary:
            raise HTTPException(status_code=404, detail="Summary not found")
        course = db.query(Course).filter(Course.id == summary.course_id, Course.user_id == user_id).first()
        return {
            "id": summary.id,
            "title": summary.title,
            "content": summary.content,
            "created_at": summary.created_at.isoformat(),
            "course_id": summary.course_id,
            "course_name": course.name if course else None,
            "course_code": course.code if course else None
        }
    finally:
        db.close()


# Flashcard endpoints
@app.post("/courses/{course_id}/flashcards")
@limiter.limit("10/minute")
async def generate_flashcards(request: Request, course_id: str, file: UploadFile = File(...), current_user: User = Depends(get_current_user)):
    """Upload study material and generate flashcards using AI."""
    logger.info(f"[DEBUG] /courses/{course_id}/flashcards request received")
    db = SessionLocal()
    try:
        user_id = current_user.id
        course = db.query(Course).filter(Course.id == course_id, Course.user_id == user_id).first()
        if not course:
            raise HTTPException(status_code=404, detail="Course not found")

        # Validate and extract text from file
        content = await validate_file_upload(file, allowed_extensions=['.pdf', '.txt', '.docx'], max_size_mb=10)
        filename = file.filename.lower()

        if filename.endswith('.pdf'):
            pdf = PyPDF2.PdfReader(io.BytesIO(content))
            text = ""
            for page in pdf.pages:
                text += (page.extract_text() or "") + "\n"
        elif filename.endswith('.txt'):
            text = content.decode('utf-8')
        else:
            raise HTTPException(status_code=400, detail="Supported formats: PDF, TXT")

        print(f"[DEBUG] Extracted {len(text)} characters from study material")

        if len(text.strip()) < 100:
            raise HTTPException(status_code=400, detail="Could not extract enough text from file")

        print(f"[DEBUG] Generating flashcards from {len(text)} characters")

        # Truncate if too long
        max_chars = 15000
        if len(text) > max_chars:
            text = text[:max_chars]

        flashcards_data = []
        api_key = os.getenv("OPENAI_API_KEY")
        if api_key:
            # Generate flashcards with OpenAI
            system_prompt = """You are a study assistant. Generate flashcards from the provided study material.

Return ONLY a valid JSON array with 10-20 flashcards:
[
    {"front": "Question or term", "back": "Answer or definition"},
    ...
]

Focus on:
- Key concepts and definitions
- Important facts and dates
- Formulas and their applications
- Cause and effect relationships
- Compare and contrast items

Make questions clear and answers concise but complete."""

            try:
                response = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": f"Generate flashcards from this material:\n\n{text}"}
                    ],
                    temperature=0.3,
                    max_tokens=4000
                )

                result = response.choices[0].message.content
                print(f"[DEBUG] Flashcard response: {result[:500]}...")

                flashcards_data = parse_json_response(result)
                print(f"[DEBUG] Generated {len(flashcards_data)} flashcards")
            except Exception as e:
                _raise_if_openai_error(e)  # don't silently fall back on auth / rate-limit
                print(f"[WARN] OpenAI flashcard generation failed: {e}")
                flashcards_data = generate_flashcards_fallback(text)
        else:
            print("[WARN] OPENAI_API_KEY not set; using fallback flashcard generator")
            flashcards_data = generate_flashcards_fallback(text)

        # Create flashcard set
        flashcard_set = FlashcardSet(
            user_id=user_id,
            course_id=course_id,
            name=file.filename.rsplit('.', 1)[0]  # Use filename without extension
        )
        db.add(flashcard_set)
        db.flush()

        # Create flashcards
        for fc in flashcards_data:
            flashcard = Flashcard(
                user_id=user_id,
                flashcard_set_id=flashcard_set.id,
                front=fc.get("front", ""),
                back=fc.get("back", "")
            )
            db.add(flashcard)

        db.commit()
        print(f"[DEBUG] Inserted {len(flashcards_data)} flashcards for course {course_id}")

        return {
            "flashcard_set": {
                "id": flashcard_set.id,
                "name": flashcard_set.name,
                "card_count": len(flashcards_data)
            },
            "flashcards": flashcards_data
        }

    except json.JSONDecodeError as e:
        print(f"[ERROR] Failed to parse flashcards JSON: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate flashcards")
    except Exception as e:
        print(f"[ERROR] Error generating flashcards: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="An error occurred while generating flashcards. Please try again.")
    finally:
        db.close()


@app.get("/flashcard-sets/{set_id}")
def get_flashcard_set(set_id: str, current_user: User = Depends(get_current_user)):
    """Get a flashcard set with all its cards."""
    db = SessionLocal()
    try:
        user_id = current_user.id
        flashcard_set = db.query(FlashcardSet).filter(FlashcardSet.id == set_id, FlashcardSet.user_id == user_id).first()
        if not flashcard_set:
            raise HTTPException(status_code=404, detail="Flashcard set not found")

        return {
            "id": flashcard_set.id,
            "name": flashcard_set.name,
            "course_id": flashcard_set.course_id,
            "flashcards": [
                {"id": fc.id, "front": fc.front, "back": fc.back}
                for fc in flashcard_set.flashcards
            ]
        }
    finally:
        db.close()


@app.delete("/flashcard-sets/{set_id}")
def delete_flashcard_set(set_id: str, current_user: User = Depends(get_current_user)):
    """Delete a flashcard set."""
    db = SessionLocal()
    try:
        user_id = current_user.id
        flashcard_set = db.query(FlashcardSet).filter(FlashcardSet.id == set_id, FlashcardSet.user_id == user_id).first()
        if not flashcard_set:
            raise HTTPException(status_code=404, detail="Flashcard set not found")

        db.delete(flashcard_set)
        db.commit()
        return {"message": "Flashcard set deleted"}
    finally:
        db.close()


@app.post("/upload")
@limiter.limit("5/minute")
async def upload_syllabus(request: Request, file: UploadFile = File(...), current_user: User = Depends(get_current_user)):
    logger.info("[DEBUG] /upload request received")
    user_id = current_user.id

    # Validate file upload — PDF and DOCX both supported
    content = await validate_file_upload(file, allowed_extensions=['.pdf', '.docx'], max_size_mb=10)
    filename = (file.filename or "").lower()

    try:
        print(f"[DEBUG] Processing file: {file.filename}")

        if filename.endswith(".docx"):
            text = extract_text_from_docx(content)
        else:
            text = extract_text_from_pdf(content)

        print(f"[DEBUG] Total text extracted: {len(text)} characters")

        if len(text.strip()) < 50:
            return {
                "course": None,
                "deadlines": [],
                "debug": {
                    "text_length": len(text),
                    "dates_found": 0,
                    "message": "Could not extract text from the uploaded file. Make sure it contains readable text."
                }
            }

        # PASS 1: Extract course metadata
        metadata = extract_course_metadata(text)

        # PASS 2: Extract deadlines with context
        deadlines_data = extract_deadlines_with_context(text, metadata)

        # Save to database
        db = SessionLocal()
        try:
            # Parse course code from name (e.g., "FINC 313 - Corporate Finance" -> "FINC 313")
            course_name = metadata.get("course_name", "Unknown Course")
            course_code = None
            if " - " in course_name:
                parts = course_name.split(" - ")
                course_code = parts[0].strip()
                course_name = parts[1].strip() if len(parts) > 1 else course_name

            # Create course
            course = Course(
                user_id=user_id,
                name=course_name,
                code=course_code,
                semester=metadata.get("semester"),
                start_date=datetime.strptime(metadata["start_date"], "%Y-%m-%d").date() if metadata.get("start_date") else None,
                end_date=datetime.strptime(metadata["end_date"], "%Y-%m-%d").date() if metadata.get("end_date") else None,
                course_info=metadata.get("course_info")
            )
            db.add(course)
            db.flush()  # Get the course ID

            # Create deadlines
            for d in deadlines_data:
                deadline = Deadline(
                    user_id=user_id,
                    course_id=course.id,
                    date=d.get("date"),
                    time=d.get("time"),
                    type=d.get("type"),
                    title=d.get("title"),
                    description=d.get("context") or d.get("description"),
                    recurring=d.get("recurring", False),
                    frequency=d.get("frequency"),
                    day_of_week=d.get("day_of_week")
                )
                db.add(deadline)

            db.commit()

            return {
                "course": {
                    "id": course.id,
                    "name": course.name,
                    "code": course.code,
                    "semester": course.semester,
                    "start_date": str(course.start_date) if course.start_date else None,
                    "end_date": str(course.end_date) if course.end_date else None,
                    "course_info": course.course_info
                },
                "deadlines": deadlines_data,
                "debug": {
                    "text_length": len(text),
                    "dates_found": len(deadlines_data),
                    "message": f"Found {len(deadlines_data)} deadlines using AI (2-pass extraction)"
                }
            }
        finally:
            db.close()

    except Exception as e:
        print(f"[ERROR] Error processing PDF: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="An error occurred while processing the PDF. Please try again.")


# ============= Quiz Endpoints =============

@app.post("/courses/{course_id}/generate-quiz")
@limiter.limit("10/minute")
async def generate_quiz(request: Request, course_id: str, file: UploadFile = File(...), current_user: User = Depends(get_current_user)):
    """Generate a multiple-choice quiz from study materials."""
    db = SessionLocal()
    try:
        user_id = current_user.id
        course = db.query(Course).filter(Course.id == course_id, Course.user_id == user_id).first()
        if not course:
            raise HTTPException(status_code=404, detail="Course not found")

        # Validate and extract text from file
        content = await validate_file_upload(file, allowed_extensions=['.pdf', '.txt', '.docx'], max_size_mb=10)
        filename = file.filename.lower()

        if filename.endswith('.pdf'):
            pdf = PyPDF2.PdfReader(io.BytesIO(content))
            text = ""
            for page in pdf.pages:
                text += (page.extract_text() or "") + "\n"
        elif filename.endswith('.txt'):
            text = content.decode('utf-8')
        elif filename.endswith('.docx'):
            # For simplicity, treat as text; proper DOCX parsing would require python-docx
            text = content.decode('utf-8', errors='ignore')
        else:
            raise HTTPException(status_code=400, detail="Supported formats: PDF, TXT, DOCX")

        print(f"[DEBUG] Extracted {len(text)} characters for quiz generation")

        if len(text.strip()) < 100:
            raise HTTPException(status_code=400, detail="Could not extract enough text from file")

        # Truncate if too long
        max_chars = 15000
        if len(text) > max_chars:
            text = text[:max_chars]

        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise HTTPException(status_code=500, detail="Quiz generation requires OpenAI API key")

        # Generate quiz with OpenAI
        system_prompt = """You are a study assistant. Generate a multiple-choice quiz from the provided study material.

Return ONLY a valid JSON object with this structure:
{
    "questions": [
        {
            "question": "Clear question text",
            "options": ["A) First option", "B) Second option", "C) Third option", "D) Fourth option"],
            "correct_answer": "B",
            "explanation": "Brief explanation of why this is correct"
        }
    ]
}

Guidelines:
- Generate 7 questions
- Each question should have exactly 4 options (A, B, C, D)
- Questions should test understanding, not just memorization
- Include a mix of difficulty levels
- Make distractors (wrong answers) plausible
- Keep explanations concise (1-2 sentences)"""

        try:
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"Generate a quiz from this material:\n\n{text}"}
                ],
                temperature=0.4,
                max_tokens=4000
            )

            result = response.choices[0].message.content
            print(f"[DEBUG] Quiz response: {result[:500]}...")

            quiz_data = parse_json_response(result)

            # Handle both array and object responses
            if isinstance(quiz_data, list):
                questions = quiz_data
            elif isinstance(quiz_data, dict) and "questions" in quiz_data:
                questions = quiz_data["questions"]
            else:
                raise ValueError("Invalid quiz format from AI")

            print(f"[DEBUG] Generated {len(questions)} quiz questions")

        except Exception as e:
            _raise_if_openai_error(e)
            print(f"[ERROR] Quiz generation failed: {e}")
            raise HTTPException(status_code=500, detail="Failed to generate quiz questions")

        # Create quiz
        quiz = Quiz(
            user_id=user_id,
            course_id=course_id,
            name=file.filename.rsplit('.', 1)[0]  # Use filename without extension
        )
        db.add(quiz)
        db.flush()

        # Create questions
        for i, q in enumerate(questions):
            options = q.get("options", [])
            if isinstance(options, list):
                options_json = json.dumps(options)
            else:
                options_json = json.dumps([])

            question = QuizQuestion(
                user_id=user_id,
                quiz_id=quiz.id,
                question=q.get("question", ""),
                options=options_json,
                correct_answer=q.get("correct_answer", "A"),
                explanation=q.get("explanation", ""),
                order_num=str(i)
            )
            db.add(question)

        db.commit()
        print(f"[DEBUG] Created quiz with {len(questions)} questions for course {course_id}")

        return {
            "quiz": {
                "id": quiz.id,
                "name": quiz.name,
                "question_count": len(questions)
            }
        }

    except HTTPException:
        raise
    except Exception as e:
        print(f"[ERROR] Error generating quiz: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="An error occurred while generating the quiz. Please try again.")
    finally:
        db.close()


@app.get("/quizzes/{quiz_id}")
def get_quiz(quiz_id: str, current_user: User = Depends(get_current_user)):
    """Get a quiz with all its questions."""
    db = SessionLocal()
    try:
        user_id = current_user.id
        quiz = db.query(Quiz).filter(Quiz.id == quiz_id, Quiz.user_id == user_id).first()
        if not quiz:
            raise HTTPException(status_code=404, detail="Quiz not found")

        questions = []
        for q in sorted(quiz.questions, key=lambda x: x.order_num):
            try:
                options = json.loads(q.options)
            except (json.JSONDecodeError, TypeError, ValueError) as e:
                print(f"[WARNING] Failed to parse options for question {q.id}: {str(e)}")
                options = []

            questions.append({
                "id": q.id,
                "question": q.question,
                "options": options,
                # Don't include correct_answer or explanation in GET - only after submission
            })

        return {
            "id": quiz.id,
            "name": quiz.name,
            "course_id": quiz.course_id,
            "question_count": len(questions),
            "questions": questions
        }
    finally:
        db.close()


@app.post("/quizzes/{quiz_id}/submit")
def submit_quiz(quiz_id: str, submission: QuizSubmission, current_user: User = Depends(get_current_user)):
    """Submit quiz answers and get results."""
    db = SessionLocal()
    try:
        user_id = current_user.id
        quiz = db.query(Quiz).filter(Quiz.id == quiz_id, Quiz.user_id == user_id).first()
        if not quiz:
            raise HTTPException(status_code=404, detail="Quiz not found")

        results = []
        correct_count = 0
        total_count = len(quiz.questions)

        for q in sorted(quiz.questions, key=lambda x: x.order_num):
            user_answer = submission.answers.get(q.id, "")
            is_correct = user_answer.upper() == q.correct_answer.upper()

            if is_correct:
                correct_count += 1

            try:
                options = json.loads(q.options)
            except (json.JSONDecodeError, TypeError, ValueError) as e:
                print(f"[WARNING] Failed to parse options for question {q.id}: {str(e)}")
                options = []

            results.append({
                "question_id": q.id,
                "question": q.question,
                "options": options,
                "user_answer": user_answer,
                "correct_answer": q.correct_answer,
                "is_correct": is_correct,
                "explanation": q.explanation
            })

        score_percentage = round((correct_count / total_count) * 100) if total_count > 0 else 0

        return {
            "quiz_id": quiz_id,
            "quiz_name": quiz.name,
            "score": correct_count,
            "total": total_count,
            "percentage": score_percentage,
            "results": results
        }
    finally:
        db.close()


@app.get("/courses/{course_id}/quizzes")
def get_course_quizzes(course_id: str, current_user: User = Depends(get_current_user)):
    """Get all quizzes for a course."""
    db = SessionLocal()
    try:
        user_id = current_user.id
        course = db.query(Course).filter(Course.id == course_id, Course.user_id == user_id).first()
        if not course:
            raise HTTPException(status_code=404, detail="Course not found")

        return [
            {
                "id": quiz.id,
                "name": quiz.name,
                "question_count": len(quiz.questions),
                "created_at": quiz.created_at.isoformat() if quiz.created_at else None
            }
            for quiz in course.quizzes
        ]
    finally:
        db.close()


@app.delete("/quizzes/{quiz_id}")
def delete_quiz(quiz_id: str, current_user: User = Depends(get_current_user)):
    """Delete a quiz."""
    db = SessionLocal()
    try:
        user_id = current_user.id
        quiz = db.query(Quiz).filter(Quiz.id == quiz_id, Quiz.user_id == user_id).first()
        if not quiz:
            raise HTTPException(status_code=404, detail="Quiz not found")

        db.delete(quiz)
        db.commit()
        return {"message": "Quiz deleted"}
    finally:
        db.close()


@app.post("/feedback")
@limiter.limit("5/minute")
def submit_feedback(request: Request, payload: FeedbackRequest, current_user: User = Depends(get_current_user)):
    """Submit user feedback. Saves to DB and sends notification email via Resend."""
    valid_types = ("bug", "feature", "general")
    if payload.feedback_type not in valid_types:
        raise HTTPException(status_code=400, detail=f"feedback_type must be one of: {', '.join(valid_types)}")

    if not payload.message.strip():
        raise HTTPException(status_code=400, detail="message cannot be empty")

    db = SessionLocal()
    try:
        feedback = Feedback(
            user_email=current_user.email,
            feedback_type=payload.feedback_type,
            message=payload.message.strip(),
        )
        db.add(feedback)
        db.commit()

        if RESEND_API_KEY and FEEDBACK_EMAIL:
            try:
                resend.api_key = RESEND_API_KEY
                type_label = {"bug": "Bug Report", "feature": "Feature Request", "general": "General Feedback"}.get(payload.feedback_type, payload.feedback_type)
                resend.Emails.send({
                    "from": "onboarding@resend.dev",
                    "to": [FEEDBACK_EMAIL],
                    "subject": f"[ClassMate Feedback] {type_label} from {current_user.email}",
                    "html": (
                        f"<h2>{type_label}</h2>"
                        f"<p><strong>From:</strong> {current_user.email}</p>"
                        f"<p><strong>Message:</strong></p>"
                        f"<p>{payload.message.strip()}</p>"
                        f"<hr><p style='color:#999;font-size:12px'>Feedback ID: {feedback.id}</p>"
                    ),
                })
            except Exception as e:
                logger.warning(f"[Feedback] Failed to send email notification: {e}")

        return {"message": "Feedback submitted successfully"}
    finally:
        db.close()
